import os
import logging
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class DocxGenerator:
    """Word生成器类
    
    负责将翻译后的文本生成Word文档，支持文本、表格和图像的生成。
    """
    
    def __init__(self):
        """初始化DocxGenerator对象"""
        logger.info("Word生成器初始化完成")
    
    def _clean_xml_compatible_text(self, text):
        """清理文本，确保它只包含XML兼容的字符
        
        Args:
            text (str): 要清理的文本
            
        Returns:
            str: 清理后的文本
        """
        if not text:
            return ""
        
        # 移除NULL字节和控制字符（保留制表符、换行符和回车符）
        cleaned_text = ""
        for char in text:
            char_code = ord(char)
            # 保留XML兼容的字符：
            # 1. 可打印的ASCII字符（32-126）
            # 2. 制表符（9）、换行符（10）、回车符（13）
            # 3. Unicode字符（128+）
            if (32 <= char_code <= 126) or char_code in (9, 10, 13) or char_code >= 128:
                cleaned_text += char
        
        return cleaned_text
    
    def generate_docx(self, translated_content, images, output_docx_path, target_lang="zh"):
        """生成翻译后的Word文档
        
        Args:
            translated_content (dict): 翻译后的内容
                - blocks (list): 每页的完整文本块列表 (翻译后的文本块)
                    - page_num (int): 页码
                    - text_blocks (list): TextBlock对象列表
                        - block_text (str): 完整文本块内容 (已翻译)
                        - block_bbox (tuple): 文本块位置 (x0, y0, x1, y1)
                        - block_no (int): 块编号
                        - block_type (int): 块类型
                        - font (str): 字体名称
                        - font_size (float): 字体大小
                        - color (int): 颜色值
                        - flags (int): 样式标记
                - merged_translations (list): 合并后的翻译结果
                    - text (str): 合并后的翻译文本
                    - page_num (int): 页码
                    - font (str): 字体名称
                    - font_size (float): 字体大小
                    - color (int): 颜色值
                    - bold (bool): 是否粗体
                    - italic (bool): 是否斜体
                    - bbox (tuple): 文本块边界框 (x0, y0, x1, y1) - 新增
                - tables (list): 翻译后的表格列表
            images (list): 图像信息列表
                - page_num (int): 页码
                - image_idx (int): 图像索引
                - image_path (str): 图像路径
                - bbox (tuple): 图像位置
            output_docx_path (str): 输出Word文件路径
            target_lang (str): 目标语言代码
        """
        logger.info(f"开始生成Word文档，输出文件: {output_docx_path}, 目标语言: {target_lang}")
        
        try:
            # 创建新的Word文档
            doc = Document()
            logger.info("创建Word文档成功")
            
            # 按页码组织合并后的翻译结果
            merged_by_page = {}
            for item in translated_content.get('merged_translations', []):
                page_num = item.page_num
                if page_num not in merged_by_page:
                    merged_by_page[page_num] = []
                merged_by_page[page_num].append(item)
            
            # 按页码组织图像
            images_by_page = {}
            for image in images:
                page_num = image.page_num
                if page_num not in images_by_page:
                    images_by_page[page_num] = []
                images_by_page[page_num].append(image)
            
            # 按页码组织表格
            tables_by_page = {}
            if 'tables' in translated_content:
                for table in translated_content['tables']:
                    page_num = table.page_num  # 使用page_num属性
                    if page_num not in tables_by_page:
                        tables_by_page[page_num] = []
                    tables_by_page[page_num].append(table)
            
            # 按页码组织原始文本块
            original_blocks_by_page = {}
            if 'blocks' in translated_content:
                for page in translated_content['blocks']:
                    page_num = page.page_num
                    if page_num not in original_blocks_by_page:
                        original_blocks_by_page[page_num] = []
                    # 原始文本块已经按垂直位置排序
                    original_blocks_by_page[page_num] = page.text_blocks
            
            # 获取所有页码的列表
            all_pages = set(merged_by_page.keys())
            all_pages.update(images_by_page.keys())
            all_pages.update(tables_by_page.keys())
            sorted_pages = sorted(all_pages)
            
            # 处理每页内容
            for i, page_num in enumerate(sorted_pages):
                logger.info(f"处理第 {page_num} 页")
                
                # 收集当前页的所有元素
                page_elements = []
                
                # 添加文本块元素
                text_blocks = merged_by_page.get(page_num, [])
                for block_idx, block in enumerate(text_blocks):
                    logger.info(f"处理第 {page_num} 页 文本块 {block_idx+1}: 内容='{block.block_text[:50]}...'")
                    
                    page_elements.append({
                        'type': 'text',
                        'content': block,
                        'block_idx': block_idx
                    })
                
                # 添加图像元素
                page_images = images_by_page.get(page_num, [])
                for image_idx, image in enumerate(page_images):
                    logger.info(f"处理第 {page_num} 页 图像 {image_idx+1}: 路径={image.image_path}")
                    
                    page_elements.append({
                        'type': 'image',
                        'content': image
                    })
                
                # 添加表格元素
                page_tables = tables_by_page.get(page_num, [])
                for table_idx, table in enumerate(page_tables):
                    logger.info(f"处理第 {page_num} 页 表格 {table_idx+1}")
                    
                    page_elements.append({
                        'type': 'table',
                        'content': table
                    })
                
                # 记录元素列表
                logger.info(f"第 {page_num} 页元素列表:")
                for i, elem in enumerate(page_elements):
                    elem_type = elem['type']
                    if elem_type == 'text':
                        content = elem['content'].block_text[:30] + '...'
                    elif elem_type == 'image':
                        content = f"图像: {elem['content'].image_path}"
                    else:
                        content = f"{elem_type}: {elem['content']}"
                    logger.info(f"  元素 {i+1}: 类型={elem_type}, 内容='{content}'")
                
                # 处理元素
                self._process_page_elements(doc, page_elements, original_blocks_by_page, page_num)
                
                # 只在不是最后一页时添加分页符
                if i < len(sorted_pages) - 1:
                    # 添加分页符
                    doc.add_page_break()
            
            # 保存文档
            output_dir = os.path.dirname(output_docx_path)
            os.makedirs(output_dir, exist_ok=True)
            
            doc.save(output_docx_path)
            logger.info(f"Word文档生成完成，输出文件: {output_docx_path}")
            
        except Exception as e:
            logger.error(f"生成Word文档时出错: {str(e)}", exc_info=True)
            raise Exception(f"生成Word文档时出错: {str(e)}")
    
    def _find_chart_position(self, image, original_blocks):
        """查找图表在原始文本块中的位置
        
        Args:
            image: 图像对象
            original_blocks: 原始文本块列表（已排序）
            
        Returns:
            tuple: (before_block, after_block) - 图表前后的原始块
        """
        chart_y = image.bbox[1]
        
        # 在已排序的原始文本块中查找图表前后的块
        # 由于原始文本块已经按垂直位置排序，可以使用线性查找
        before_block = None
        after_block = None
        
        for i, block in enumerate(original_blocks):
            block_y = block.block_bbox[1]
            if block_y <= chart_y:
                before_block = block
            else:
                after_block = block
                break
        
        logger.info(f"图表位置查找完成: before_block={before_block.block_no if before_block else None}, after_block={after_block.block_no if after_block else None}")
        return before_block, after_block
    
    def _find_merged_block(self, before_block, after_block, merged_blocks, chart_y=None):
        """查找包含图表前后原始块的合并块
        
        Args:
            before_block: 图表前的原始块
            after_block: 图表后的原始块
            merged_blocks: 合并块列表
            chart_y: 图表的y坐标（可选），用于计算在合并块内的插入点
            
        Returns:
            tuple: (merged_block, position, is_within_merged_block, insertion_ratio)
                - merged_block: 包含图表的合并块
                - position: 位置关系 ('before', 'after', 'within')
                - is_within_merged_block: 图表是否位于合并块内部
                - insertion_ratio: 图表在合并块中的相对位置比例（0-1）
        """
        for block in merged_blocks:
            original_blocks = block.original_blocks
            original_block_nos = [b.block_no for b in original_blocks]
            
            # 检查before_block和after_block是否都在同一个合并块中
            before_in_block = before_block and before_block.block_no in original_block_nos
            after_in_block = after_block and after_block.block_no in original_block_nos
            
            if before_in_block and after_in_block:
                # 图表位于合并块内部
                logger.info(f"图表位于合并块内部: {block.block_text[:30]}...")
                
                if chart_y is not None:
                    # 计算插入点比例
                    merged_start_y = min(b.block_bbox[1] for b in original_blocks)
                    merged_end_y = max(b.block_bbox[3] for b in original_blocks)
                    merged_height = merged_end_y - merged_start_y
                    
                    if merged_height > 0:
                        insertion_ratio = (chart_y - merged_start_y) / merged_height
                        insertion_ratio = max(0.1, min(0.9, insertion_ratio))  # 限制在10%-90%之间
                        logger.info(f"计算合并块内插入点比例: {insertion_ratio:.2f}, 合并块范围: y0={merged_start_y}, y1={merged_end_y}")
                        return block, 'within', True, insertion_ratio
                
                return block, 'within', True, 0.5  # 默认中间位置
            
            # 检查before_block是否在当前合并块中（使用块编号）
            if before_block and before_block.block_no in original_block_nos:
                logger.info(f"找到包含before_block的合并块: {block.block_text[:30]}...")
                return block, 'after', False, None  # 图表在before_block之后
            
            # 检查after_block是否在当前合并块中（使用块编号）
            if after_block and after_block.block_no in original_block_nos:
                logger.info(f"找到包含after_block的合并块: {block.block_text[:30]}...")
                return block, 'before', False, None  # 图表在after_block之前
        
        logger.info("未找到包含图表前后原始块的合并块")
        return None, None, False, None
    
    def _process_page_elements(self, doc, page_elements, original_blocks_by_page, page_num):
        """处理页面元素，按顺序添加到Word文档
        
        Args:
            doc: Word文档对象
            page_elements (list): 页面元素列表
            original_blocks_by_page: 按页码组织的原始文本块
            page_num: 当前页码
        """
        logger.info(f"开始处理页面元素，元素数量: {len(page_elements)}")
        
        # 跟踪已处理的文本块索引
        processed_blocks = set()
        
        # 获取当前页面的原始文本块（已排序）
        original_blocks = original_blocks_by_page.get(page_num, [])
        
        # 分离文本块和图表元素
        text_elements = []
        chart_elements = []
        
        for element in page_elements:
            if element['type'] == 'text':
                text_elements.append(element)
            else:
                chart_elements.append(element)
        
        # 处理图表元素，确定它们的插入位置
        chart_insertions = []
        for chart in chart_elements:
            if chart['type'] == 'image' or chart['type'] == 'table':
                # 查找图表/表格在原始文本块中的位置
                before_block, after_block = None, None
                chart_y = None
                if hasattr(chart['content'], 'bbox'):
                    chart_y = chart['content'].bbox[1]
                    # 如果有bbox属性，使用与图像相同的位置查找逻辑
                    before_block, after_block = self._find_chart_position(chart['content'], original_blocks)
                
                # 查找包含图表/表格前后原始块的合并块
                merged_blocks = [e['content'] for e in text_elements]
                merged_block, position, is_within, insertion_ratio = self._find_merged_block(before_block, after_block, merged_blocks, chart_y)
                
                if merged_block:
                    chart_insertions.append((merged_block, position, chart, insertion_ratio))
                else:
                    # 如果没有找到合适的合并块，直接添加图表/表格
                    chart_insertions.append((None, 'end', chart, None))
        
        # 处理文本块和图表
        for element in text_elements:
            block = element['content']
            block_idx = element['block_idx']
            
            if block_idx not in processed_blocks:
                # 检查是否有图表需要插入到当前块之前
                for merged_block, position, chart, insertion_ratio in chart_insertions:
                    if merged_block == block and position == 'before':
                        # 插入图表到当前块之前
                        if chart['type'] == 'image':
                            logger.info(f"  在文本块之前插入图像: {chart['content'].image_path}")
                            self._add_image(doc, chart['content'])
                        elif chart['type'] == 'table':
                            logger.info(f"  在文本块之前插入表格")
                            self._add_table(doc, chart['content'])
                
                # 检查是否有图表需要插入到当前块内部
                for merged_block, position, chart, insertion_ratio in chart_insertions:
                    if merged_block == block and position == 'within':
                        # 在合并块内部插入图表，拆分文本
                        ratio_str = f"{insertion_ratio:.2f}" if insertion_ratio is not None else "0.5"
                        logger.info(f"  在合并块内部插入图表，插入点比例: {ratio_str}")
                        self._insert_element_in_text_block(doc, block, chart, insertion_ratio)
                        processed_blocks.add(block_idx)
                        logger.info(f"  合并块内部插入完成")
                        break
                
                # 如果没有在块内部插入，则正常添加文本块
                if block_idx not in processed_blocks:
                    # 添加文本块
                    logger.info(f"  添加文本块 {block_idx+1}: 内容='{block.block_text[:30]}...'")
                    self._add_merged_text(doc, block)
                    processed_blocks.add(block_idx)
                    logger.info(f"  文本块 {block_idx+1} 处理完成")
                    
                    # 检查是否有图表需要插入到当前块之后
                    for merged_block, position, chart, insertion_ratio in chart_insertions:
                        if merged_block == block and position == 'after':
                            # 插入图表到当前块之后
                            if chart['type'] == 'image':
                                logger.info(f"  在文本块之后插入图像: {chart['content'].image_path}")
                                self._add_image(doc, chart['content'])
                            elif chart['type'] == 'table':
                                logger.info(f"  在文本块之后插入表格")
                                self._add_table(doc, chart['content'])
        
        # 处理需要添加到文档末尾的图表
        for merged_block, position, chart, insertion_ratio in chart_insertions:
            if position == 'end':
                if chart['type'] == 'image':
                    logger.info(f"  在文档末尾插入图像: {chart['content'].image_path}")
                    self._add_image(doc, chart['content'])
                elif chart['type'] == 'table':
                    logger.info(f"  在文档末尾插入表格")
                    self._add_table(doc, chart['content'])
        
        logger.info(f"页面元素处理完成，已处理 {len(processed_blocks)} 个文本块")
    
    def _insert_element_in_text_block(self, doc, text_block, element, insertion_point):
        """在文本块内部插入元素，拆分文本
        
        Args:
            doc: Word文档对象
            text_block: 文本块对象
            element: 要插入的元素
            insertion_point: 插入点比例（0-1）
        """
        logger.info(f"拆分合并块，插入点: {insertion_point:.2f}, 合并块内容: '{text_block.block_text[:50]}...'")
        
        text = text_block.block_text
        text_length = len(text)
        
        # 计算插入位置
        split_index = int(text_length * insertion_point)
        
        # 尝试在单词边界拆分
        if split_index > 0 and split_index < text_length:
            # 向前查找空格或标点
            while split_index > 0 and text[split_index-1].isalnum():
                split_index -= 1
        
        # 拆分为前后两段
        text_before = text[:split_index].rstrip()
        text_after = text[split_index:].lstrip()
        
        logger.info(f"拆分结果: 前半部分='{text_before[:30]}...', 后半部分='{text_after[:30]}...'")
        
        # 添加前段文本
        if text_before:
            from models.merged_block import MergedBlock
            before_block = MergedBlock(
                block_text=text_before,
                original_blocks=text_block.original_blocks,
                max_width=text_block.max_width,
                max_height=text_block.max_height
            )
            self._add_merged_text(doc, before_block)
            logger.info("添加前段文本完成")
        
        # 添加元素
        if element['type'] == 'image':
            self._add_image(doc, element['content'])
            logger.info("添加图像完成")
        elif element['type'] == 'table':
            self._add_table(doc, element['content'])
            logger.info("添加表格完成")
        
        # 添加后段文本
        if text_after:
            from models.merged_block import MergedBlock
            after_block = MergedBlock(
                block_text=text_after,
                original_blocks=text_block.original_blocks,
                max_width=text_block.max_width,
                max_height=text_block.max_height
            )
            self._add_merged_text(doc, after_block)
            logger.info("添加后段文本完成")
    
    def _add_merged_text(self, doc, merged_item):
        """添加合并后的文本到Word文档
        
        Args:
            doc: Word文档对象
            merged_item: 合并后的文本项（MergedBlock对象）
                - block_text (str): 合并后的翻译文本
                - font (str): 字体名称
                - font_size (float): 字体大小
                - color (int): 颜色值
                - bold (bool): 是否粗体
                - italic (bool): 是否斜体
        """
        # 创建段落
        paragraph = doc.add_paragraph()
        
        # 清理文本，确保它只包含XML兼容的字符
        cleaned_text = self._clean_xml_compatible_text(merged_item.block_text)
        
        # 添加文本
        run = paragraph.add_run(cleaned_text)
        
        # 应用样式
        try:
            # 设置字体
            run.font.name = merged_item.font
            
            # 设置字体大小
            font_size = merged_item.font_size
            run.font.size = Pt(font_size)
            logger.info(f"应用字体大小: {font_size} 到文本: '{cleaned_text[:50]}...'")
            
            # 设置颜色
            color = merged_item.color
            if color > 0xFFFFFF:  # ARGB
                r = (color >> 16) & 0xFF
                g = (color >> 8) & 0xFF
                b = color & 0xFF
            else:  # RGB
                r = (color >> 16) & 0xFF
                g = (color >> 8) & 0xFF
                b = color & 0xFF
            run.font.color.rgb = RGBColor(r, g, b)
            
            # 设置粗体和斜体
            run.bold = merged_item.bold
            run.italic = merged_item.italic
            
            logger.debug(f"添加合并文本，内容: '{cleaned_text[:50]}...'，字体: {merged_item.font}，大小: {font_size}")
        except Exception as e:
            logger.warning(f"设置合并文本样式失败: {e}")
    
    def _add_text_block(self, doc, text_block):
        """添加文本块到Word文档
        
        Args:
            doc: Word文档对象
            text_block: 文本块对象
        """
        # 创建段落
        paragraph = doc.add_paragraph()
        
        # 清理文本，确保它只包含XML兼容的字符
        cleaned_text = self._clean_xml_compatible_text(text_block.block_text)
        
        # 添加文本
        run = paragraph.add_run(cleaned_text)
        
        # 应用样式
        self._apply_style(run, text_block)
        
        logger.debug(f"添加文本块，内容: '{cleaned_text[:50]}...'，字体: {text_block.font}，大小: {text_block.font_size}")
    
    def _apply_style(self, run, text_block):
        """应用样式到文本
        
        Args:
            run: Word文本运行对象
            text_block: 文本块对象
        """
        # 设置字体
        try:
            run.font.name = text_block.font
        except Exception as e:
            logger.warning(f"设置字体失败: {e}，使用默认字体")
        
        # 设置字体大小
        try:
            run.font.size = Pt(text_block.font_size)
        except Exception as e:
            logger.warning(f"设置字体大小失败: {e}")
        
        # 设置颜色
        try:
            color = text_block.color
            if color > 0xFFFFFF:  # ARGB
                r = (color >> 16) & 0xFF
                g = (color >> 8) & 0xFF
                b = color & 0xFF
            else:  # RGB
                r = (color >> 16) & 0xFF
                g = (color >> 8) & 0xFF
                b = color & 0xFF
            run.font.color.rgb = RGBColor(r, g, b)
        except Exception as e:
            logger.warning(f"设置颜色失败: {e}")
        
        # 设置粗体和斜体
        run.bold = text_block.bold
        run.italic = text_block.italic
    
    def _add_image(self, doc, image):
        """添加图像到Word文档
        
        Args:
            doc: Word文档对象
            image: 图像对象
        """
        image_path = image.image_path
        if os.path.exists(image_path):
            try:
                # 计算图像大小（简化处理）
                width = image.bbox[2] - image.bbox[0]
                height = image.bbox[3] - image.bbox[1]
                
                # 转换为英寸（1点 = 1/72英寸）
                width_inch = width / 72
                height_inch = height / 72
                
                # 限制最大尺寸
                max_width = 6.0  # 6英寸
                if width_inch > max_width:
                    ratio = max_width / width_inch
                    width_inch = max_width
                    height_inch = height_inch * ratio
                
                # 添加图像
                doc.add_picture(image_path, width=Inches(width_inch))
                logger.info(f"添加图像成功: {image_path}，尺寸: {width_inch:.2f}x{height_inch:.2f}英寸")
            except Exception as e:
                logger.error(f"添加图像失败: {e}")
        else:
            logger.warning(f"图像文件不存在: {image_path}")
    
    def _add_table(self, doc, table):
        """添加表格到Word文档

        Args:
            doc: Word文档对象
            table: PdfTable对象
        """
        # 使用cells属性
        table_data = table.cells
        
        if not table_data:
            logger.warning("表格数据为空，跳过")
            return
        
        # 创建表格
        num_rows = len(table_data)
        num_cols = len(table_data[0]) if num_rows > 0 else 0
        
        if num_rows > 0 and num_cols > 0:
            word_table = doc.add_table(rows=num_rows, cols=num_cols)
            
            # 填充表格数据
            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    # 使用text属性
                    cell_text = cell.text
                    # 清理文本，确保它只包含XML兼容的字符
                    cleaned_text = self._clean_xml_compatible_text(str(cell_text))
                    word_table.cell(i, j).text = cleaned_text
            
            logger.info(f"添加表格成功，{num_rows}行{num_cols}列")
        else:
            logger.warning("表格数据格式不正确，跳过")